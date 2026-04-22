import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.helpers import sanitize_filename


def generate_letter(assignment, output_dir):
    """Generate a Medical Records request letter (.docx) for a single assignment.

    Parameters
    ----------
    assignment : dict
        Must contain keys: patient_name, dos, collector.
        Optional keys: claim_number, insurance_name, provider_name, npi, address.
    output_dir : str
        Base output directory.  A sub-folder named after the collector is
        created automatically.

    Returns
    -------
    str
        Absolute path of the saved .docx file.
    """
    patient_name = str(assignment.get("patient_name", "Unknown Patient")).strip()
    dos = str(assignment.get("dos", "Unknown_Date")).strip()
    collector = str(assignment.get("collector", "Unknown_Collector")).strip()
    claim_number = assignment.get("claim_number", "") or "N/A"
    insurance_name = assignment.get("insurance_name", "") or "N/A"
    provider_name = assignment.get("provider_name", "") or "N/A"
    npi = assignment.get("npi", "") or "N/A"
    address = assignment.get("address", "") or "N/A"

    # Create collector sub-folder
    collector_dir = os.path.join(output_dir, sanitize_filename(collector))
    os.makedirs(collector_dir, exist_ok=True)

    doc = _build_document(
        patient_name=patient_name,
        dos=dos,
        collector=collector,
        claim_number=claim_number,
        insurance_name=insurance_name,
        provider_name=provider_name,
        npi=npi,
        address=address,
    )

    # Build filename: PatientName_DOS.docx  (sanitised for Windows)
    safe_patient = sanitize_filename(patient_name)
    safe_dos = sanitize_filename(dos)
    filename = f"{safe_patient}_{safe_dos}.docx"
    output_path = os.path.join(collector_dir, filename)

    doc.save(output_path)
    return output_path


def _build_document(
    patient_name,
    dos,
    collector,
    claim_number,
    insurance_name,
    provider_name,
    npi,
    address,
):
    """Build and return a python-docx Document for an MR letter."""
    doc = Document()

    # ----- Header / Title -----
    title = doc.add_heading("Medical Records Request", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_color(title.runs[0], 0x1F, 0x49, 0x7D)  # dark-blue

    doc.add_paragraph("")  # spacer

    # ----- Date line -----
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run(f"Date: {datetime.today().strftime('%B %d, %Y')}")
    run.italic = True

    doc.add_paragraph("")

    # ----- Recipient block -----
    _add_bold_field(doc, "To:", "Medical Records Department")
    if provider_name and provider_name != "N/A":
        _add_bold_field(doc, "Provider / Facility:", provider_name)
    if npi and npi != "N/A":
        _add_bold_field(doc, "NPI:", npi)
    if address and address != "N/A":
        _add_bold_field(doc, "Address:", address)

    doc.add_paragraph("")

    # ----- Subject line -----
    subj = doc.add_paragraph()
    subj_run = subj.add_run(
        f"RE: Medical Records Request — {patient_name} | DOS: {dos}"
    )
    subj_run.bold = True
    subj_run.underline = True

    doc.add_paragraph("")

    # ----- Salutation -----
    doc.add_paragraph("To Whom It May Concern,")
    doc.add_paragraph("")

    # ----- Body -----
    body = (
        f"We are writing on behalf of the assigned collector, {collector}, "
        f"to formally request copies of the medical records for the following patient:"
    )
    doc.add_paragraph(body)
    doc.add_paragraph("")

    # Patient detail table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Details"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows_data = [
        ("Patient Name", patient_name),
        ("Date of Service (DOS)", dos),
        ("Claim Number", claim_number),
        ("Insurance", insurance_name),
        ("Assigned Collector", collector),
    ]
    for label, value in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.add_paragraph("")

    # ----- Closing -----
    closing = (
        "Please provide all relevant clinical notes, diagnostic reports, "
        "treatment records, and any other documentation pertaining to the above date "
        "of service at your earliest convenience. "
        "If you have any questions or require further information, "
        "please do not hesitate to contact us."
    )
    doc.add_paragraph(closing)
    doc.add_paragraph("")
    doc.add_paragraph("Thank you for your prompt attention to this matter.")
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("")
    _add_bold_field(doc, "Collector:", collector)

    return doc


def _add_bold_field(doc, label, value):
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label} ")
    run_label.bold = True
    para.add_run(value)


def _set_run_color(run, r, g, b):
    run.font.color.rgb = RGBColor(r, g, b)
